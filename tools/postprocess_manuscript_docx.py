#!/usr/bin/env python3
"""Apply submission-safe page and front-matter fixes to the Pandoc DOCX."""

from pathlib import Path
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def insert_front_matter(document: Document) -> None:
    abstract = next((p for p in document.paragraphs if p.text.strip() == "Abstract"), None)
    if abstract is None:
        raise RuntimeError("Could not locate the Abstract heading")

    front_lines = [
        "School of Art and Design, Henan University of Engineering, No.1 Xianghe Road, Longhu, Xinzheng, Zhengzhou, Henan 451191, China",
        "Department of Light Engineering, Qilu University of Technology, No.3501 Daxue Road, Changqing District, Jinan, Shandong 250353, China",
        "Corresponding author: Zihang Xue. Contact details are maintained in the journal submission system.",
    ]
    for text in front_lines:
        paragraph = abstract.insert_paragraph_before(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(9)


def format_keywords(document: Document) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text.startswith("Visual complexity") and "Reproducibility" in text:
            paragraph.text = (
                "Keywords: Visual complexity, Synthetic stress test, Image analysis, "
                "Interpretable features, Negative results, Reproducibility"
            )
            break


def number_captions(document: Document) -> None:
    """Restore visible figure/table numbers that Pandoc omits from LaTeX captions."""
    counters = {"Image Caption": 0, "Table Caption": 0}
    labels = {"Image Caption": "Figure", "Table Caption": "Table"}
    for paragraph in document.paragraphs:
        style = paragraph.style.name
        if style not in counters:
            continue
        counters[style] += 1
        prefix = f"{labels[style]} {counters[style]}. "
        if paragraph.text.lstrip().startswith(prefix):
            continue
        run = paragraph.add_run(prefix)
        run.bold = True
        paragraph._p.remove(run._r)
        paragraph._p.insert(1 if paragraph._p.pPr is not None else 0, run._r)


def format_pages(document: Document) -> None:
    for section in document.sections:
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def remove_pandoc_alignment_markers(document: Document) -> None:
    """Remove LaTeX alignment ampersands that Pandoc exposes in OMML."""
    for text_node in document.element.xpath(".//*[local-name()='t' and text()='&']"):
        math_run = text_node.getparent()
        parent = math_run.getparent()
        parent.remove(math_run)


def format_tables(document: Document) -> None:
    """Keep wide evidence tables readable and prevent rows from clipping."""
    for table in document.tables:
        table.autofit = True
        layout = table._tbl.tblPr.first_child_found_in("w:tblLayout")
        if layout is None:
            layout = OxmlElement("w:tblLayout")
            table._tbl.tblPr.append(layout)
        layout.set(qn("w:type"), "autofit")
        for row in table.rows:
            for cell in row.cells:
                cell_width = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
                if cell_width is not None:
                    cell_width.set(qn("w:type"), "auto")
                    cell_width.set(qn("w:w"), "0")
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
                    for run in paragraph.runs:
                        run.font.size = Pt(8)


def insert_references_heading(document: Document) -> None:
    """Pandoc emits bibliography entries but may omit the visible heading."""
    first_bibliography = next(
        (p for p in document.paragraphs if p.style.name == "Bibliography"),
        None,
    )
    if first_bibliography is None:
        raise RuntimeError("Could not locate the bibliography")
    previous = first_bibliography._p.getprevious()
    if previous is not None:
        previous_text = "".join(previous.itertext()).strip()
        if previous_text == "References":
            return
    heading = first_bibliography.insert_paragraph_before("References")
    heading.style = document.styles["Heading 1"]


def format_bibliography(document: Document) -> None:
    """Keep the reference list compact enough to avoid a nearly empty last page."""
    for paragraph in document.paragraphs:
        if paragraph.style.name != "Bibliography":
            continue
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.size = Pt(9)


def main() -> None:
    path = Path(sys.argv[1]) if len(sys.argv) > 1 else Path("main_no_human.docx")
    document = Document(path)
    insert_front_matter(document)
    format_keywords(document)
    number_captions(document)
    format_pages(document)
    remove_pandoc_alignment_markers(document)
    format_tables(document)
    insert_references_heading(document)
    format_bibliography(document)
    document.save(path)


if __name__ == "__main__":
    main()
